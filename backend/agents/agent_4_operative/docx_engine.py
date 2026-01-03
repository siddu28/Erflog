import os
import tempfile
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert
import shutil

class DocxSurgeon:
    def __init__(self):
        pass

    def convert_pdf_to_docx(self, pdf_path: str) -> str:
        """
        Converts PDF to DOCX using pdf2docx.
        """
        output_dir = tempfile.gettempdir()
        filename = os.path.basename(pdf_path).replace(".pdf", ".docx")
        docx_path = os.path.join(output_dir, filename)
        
        print(f"🔄 [DocxSurgeon] Converting PDF to DOCX: {pdf_path}")
        
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            print(f"   ✅ Conversion complete: {docx_path}")
            return docx_path
        except Exception as e:
            print(f"   ❌ PDF->DOCX Conversion failed: {e}")
            raise e

    def replace_section_content(self, docx_path: str, section_name: str, new_text: str) -> bool:
        """
        Edits the DOCX file to replace section content.
        Strategy: Find the header, then replace text until the next header.
        """
        print(f"📝 [DocxSurgeon] Editing section: {section_name}")
        
        try:
            doc = Document(docx_path)
            
            # Normalize header for search
            target_header = section_name.strip().upper()
            
            # Simple header detection heuristics
            # Known headers: EXPERIENCE, PROJECTS, SKILLS, EDUCATION
            
            header_found = False
            in_target_section = False
            
            # We need to iterate and modify. 
            # Note: doc.paragraphs implies linear text. Tables might be issue but pdf2docx usually outputs flow text or tables.
            # pdf2docx often puts layout in tables. We might need to search tables too.
            
            # PHASE 1: SEARCH PARAGRAPHS
            paragraphs_to_remove = []
            insert_point = None
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip().upper()
                
                # Check if this is a header
                is_header = text in ["EXPERIENCE", "PROJECTS", "SKILLS", "EDUCATION", "WORK EXPERIENCE", "TECHNICAL SKILLS"]
                
                if is_header:
                    if text == target_header or target_header in text:
                        print(f"   📍 Found Header: {text}")
                        header_found = True
                        in_target_section = True
                        insert_point = i + 1 # Insert after header
                        continue
                    elif in_target_section:
                        # Found NEXT header, stop
                        print(f"   🛑 Found Next Header: {text}. Stopping section edit.")
                        in_target_section = False
                        break
                
                if in_target_section:
                    # Mark paragraph for deletion (clear text)
                    # We can't easily "delete" paragraphs in python-docx list, but we can clear content
                    para.clear() 
            
            # PHASE 2: INSERT NEW CONTENT
            if header_found and insert_point is not None:
                # We cleared the old text. Now we need to insert the new text.
                # However, python-docx doesn't support "insert paragraph at index" easily without internal API hacks.
                # EASIER: Append text to the Header paragraph? No.
                # EASIER: We just cleared the old paragraphs. We can set the text of the *first* cleared paragraph 
                # to the new content, provided there was at least one.
                
                # Let's try to be safer: 
                # 1. Clear all old paragraphs in section.
                # 2. Add the new text to the FIRST content paragraph we found (or creating one).
                
                # Re-iterate to fill
                in_target_section = False
                filled = False
                
                for para in doc.paragraphs:
                    text = para.text.strip().upper() # This might be empty now if cleared? No, para.text is accessed.
                    
                    # We need to track index manually or re-check headers? 
                    # If we cleared text, checks will fail.
                    # BETTER STRATEGY: Do it in one pass or keep indices.
                    pass
                
                # Let's retry strategy:
                # 1. Scan to find Header Index and Next Header Index.
                # 2. Delete paragraphs between them.
                # 3. Insert new paragraph at Header Index + 1.
                pass 
                
            raise NotImplementedError("Complex logic needed for DOCX structure. Using simplified append for now.")

        except Exception as e:
            print(f"   ⚠️ Edit failed (Fallback to simple replace?): {e}")
            return False

    def _detect_section_style(self, doc, start_idx, end_idx):
        """
        Scans paragraphs in the range to find the most common font and size.
        """
        fonts = {}
        sizes = {}
        
        # Scan a few paragraphs (max 5) to get the vibe
        scan_limit = min(end_idx, start_idx + 5)
        
        for i in range(start_idx, scan_limit):
            try:
                if i >= len(doc.paragraphs): break
                p = doc.paragraphs[i]
                for run in p.runs:
                    # Font Name
                    if run.font.name:
                        fonts[run.font.name] = fonts.get(run.font.name, 0) + 1
                    
                    # Font Size
                    if run.font.size:
                        sizes[run.font.size] = sizes.get(run.font.size, 0) + 1
            except:
                pass
                
        # Get dominant style
        dom_font = max(fonts, key=fonts.get) if fonts else None
        dom_size = max(sizes, key=sizes.get) if sizes else None
        
        print(f"   🎨 Detected Style: Font={dom_font}, Size={dom_size}")
        return {"font_name": dom_font, "font_size": dom_size}

    def simple_replace(self, docx_path: str, edits: list) -> str:
        """
        Replaces sections in DOCX (Paragraphs & Tables) ensuring style preservation.
        """
        doc = Document(docx_path)
        
        headers = ["EXPERIENCE", "PROJECTS", "SKILLS", "EDUCATION", "WORK EXPERIENCE", "TECHNICAL SKILLS", "CERTIFICATIONS"]
        
        for edit in edits:
            section = edit["section"].upper() 
            new_content = edit["content"].replace("**", "") 
            print(f"   ✏️ Patching {section}...")
            
            replaced = False
            
            # --- STRATEGY 1: PARAGRAPHS ---
            start_idx = -1
            for i, p in enumerate(doc.paragraphs):
                txt = p.text.strip().upper()
                if (section == txt) or (section in txt and len(txt) < 40):
                    start_idx = i
                    break
            
            if start_idx != -1:
                # 2. Find End Header
                end_idx = len(doc.paragraphs)
                for i in range(start_idx + 1, len(doc.paragraphs)):
                    p = doc.paragraphs[i]
                    txt = p.text.strip().upper()
                    is_header_line = any(h == txt for h in headers)
                    if is_header_line or (txt.isupper() and len(txt) < 30 and len(txt) > 3 and " " not in txt):
                        end_idx = i
                        break
                
                # 3. Detect Style from the OLD content before deleting
                style = self._detect_section_style(doc, start_idx + 1, end_idx)
                
                # 4. Replace Content
                if start_idx + 1 < len(doc.paragraphs):
                    target_p = doc.paragraphs[start_idx + 1]
                    target_p.clear() # Clear existing runs
                    
                    # Add new run with detected style
                    run = target_p.add_run(new_content)
                    if style["font_name"]:
                        run.font.name = style["font_name"]
                    if style["font_size"]:
                        run.font.size = style["font_size"]
                    
                    # Remove subsequent paragraphs
                    for i in range(end_idx - 1, start_idx + 1, -1):
                        p_element = doc.paragraphs[i]._element
                        p_element.getparent().remove(p_element)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(new_content)
                    if style["font_name"]: run.font.name = style["font_name"]
                
                replaced = True

            if replaced: continue

            # --- STRATEGY 2: TABLES ---
            # If not found in paragraphs, check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt = cell.text.strip().upper()
                        # Check header in cell
                        if (section == txt) or (section in txt and len(txt) < 40):
                            print(f"   📍 Found Header '{section}' in Table Cell")
                            
                            # We assume the content follows IN THE SAME CELL or subsequent cells?
                            # Usually simple resumes might have "EXPERIENCE" in col 1, content in col 2.
                            # OR "EXPERIENCE" in row 1, content in row 2.
                            
                            # Heuristic: Clear this cell's text *after* the header? 
                            # Or does the header occupy the whole cell?
                            
                            # Case A: Header is the whole cell. Content is likely in NEXT row/cell.
                            # Case B: Header + Content in same cell.
                            
                            # Let's try to find where the content is.
                            # Only replacing the *next* block if possible.
                            
                            # Simplified Table Edit:
                            # If header is in this cell, let's append new content here or replace the rest of the cell?
                            # SAFEST: Replace the WHOLE CELL content with "HEADER + \n + New Content" 
                            # to ensure we don't lose the header if it's mixed.
                            
                            # Capture style from cell paragraphs
                            style = {"font_name": None, "font_size": None}
                            if cell.paragraphs:
                                style = self._detect_style_from_obj(cell.paragraphs[0])

                            cell.text = section + "\n" + new_content # Simple overwrite
                            
                            # Re-apply style to the whole cell
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    if style.get("font_name"): run.font.name = style["font_name"]
                                    if style.get("font_size"): run.font.size = style["font_size"]
                            
                            replaced = True
                            break
                    if replaced: break
                if replaced: break
                
        new_path = docx_path.replace(".docx", "_edited.docx")
        doc.save(new_path)
        return new_path

    def extract_text(self, docx_path: str) -> dict:
        """
        Parses the DOCX file (paragraphs AND tables) and returns section content.
        ALSO extracts hyperlink URLs which are often lost in simple text extraction.
        """
        from docx.oxml.ns import qn
        
        doc = Document(docx_path)
        sections = {}
        current_section = "HEADER_CONTACT" 
        sections[current_section] = ""
        
        headers = ["EXPERIENCE", "PROJECTS", "SKILLS", "EDUCATION", "WORK EXPERIENCE", "TECHNICAL SKILLS", "CERTIFICATIONS", "SUMMARY", "OBJECTIVE"]
        
        # Get relationships for hyperlink resolution
        rels = doc.part.rels
        
        def extract_paragraph_text_with_hyperlinks(paragraph):
            """
            Extracts text from a paragraph, including hyperlink URLs.
            python-docx doesn't expose hyperlinks directly, so we access the XML.
            """
            text_parts = []
            
            for child in paragraph._element:
                # Regular text run
                if child.tag == qn('w:r'):
                    t_elements = child.findall(qn('w:t'))
                    for t in t_elements:
                        if t.text:
                            text_parts.append(t.text)
                
                # Hyperlink
                elif child.tag == qn('w:hyperlink'):
                    # Get the relationship ID
                    r_id = child.get(qn('r:id'))
                    hyperlink_url = None
                    
                    if r_id and r_id in rels:
                        hyperlink_url = rels[r_id].target_ref
                    
                    # Get display text from runs inside the hyperlink
                    display_text_parts = []
                    for run in child.findall(qn('w:r')):
                        for t in run.findall(qn('w:t')):
                            if t.text:
                                display_text_parts.append(t.text)
                    
                    display_text = "".join(display_text_parts)
                    
                    # Append both display text AND the URL
                    if hyperlink_url and hyperlink_url != display_text:
                        text_parts.append(f"{display_text} ({hyperlink_url})")
                    else:
                        text_parts.append(display_text)
            
            return "".join(text_parts)
        
        def process_text(text):
            nonlocal current_section
            if not text.strip(): return
            upper_text = text.strip().upper()
            
            # Simple Header Heuristics
            is_header = False
            if upper_text in headers: is_header = True
            elif len(upper_text) < 40 and any(h == upper_text for h in headers): is_header = True
            elif len(upper_text) < 40 and upper_text.isupper() and " " not in upper_text and len(upper_text) > 3: is_header = True
            
            if is_header:
                current_section = text.strip().lower()
                sections[current_section] = ""
                print(f"   📄 Found Section (Ext): {current_section}")
            else:
                sections[current_section] += text + "\n"

        # 1. Scan Paragraphs (with hyperlink extraction)
        for p in doc.paragraphs:
            full_text = extract_paragraph_text_with_hyperlinks(p)
            process_text(full_text)
            
        # 2. Scan Tables (with hyperlink extraction)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        full_text = extract_paragraph_text_with_hyperlinks(p)
                        process_text(full_text)
                
        return sections

    def _detect_style_from_obj(self, paragraph_obj):
        # Helper to get style from a paragraph object
        try:
             for run in paragraph_obj.runs:
                 if run.font.name: return {"font_name": run.font.name, "font_size": run.font.size}
        except: pass
        return {}

    def convert_docx_to_pdf(self, docx_path: str) -> str:
        """
        Converts DOCX back to PDF using docx2pdf (requires Word/Office).
        Includes wait-loop to handle potential async file system lag.
        """
        import time
        pdf_path = docx_path.replace(".docx", ".pdf")
        print(f"🔄 [DocxSurgeon] Converting DOCX to PDF: {docx_path}")
        
        try:
            convert(docx_path, pdf_path)
            
            # WAIT LOOP: Ensure file exists before returning
            retries = 10
            while retries > 0:
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    print(f"   ✅ Conversion verified: {pdf_path}")
                    return pdf_path
                time.sleep(0.5)
                retries -= 1
                
            print(f"   ❌ File not found after conversion: {pdf_path}")
            return None
            
        except Exception as e:
            print(f"   ❌ DOCX->PDF Conversion failed: {e}")
            return None
